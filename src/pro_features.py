"""Professional-tier features for Super Avvocato (V7.11).

Four tools aimed at practising lawyers, each implemented as a thin wrapper
around the existing LLM backend + BM25 corpus. Keeping them out of
``brain.py`` (which already owns the citizen-facing Q&A pipeline) so the
mental model stays clean:

    ① stress_test_hearing   — red-team pre-udienza (opponent counsel sim)
    ②  audit_citations       — verify every "Neni X i Kodit Y" in a text
    ③ draft_act             — turn NL case description into filing-ready act
    ④ compute_deadline_cascade — pure rules engine over KPC/KPP deadlines

Each public function returns a plain ``dict`` ready for JSON serialisation;
persistence + HTTP wiring live in ``web.py``. No global state — callers
pass the backend + index they already have loaded.
"""
from __future__ import annotations

import json
import re
import textwrap
from dataclasses import asdict
from datetime import date, datetime, timedelta, timezone
from pathlib import Path

from .backends import LLMBackend
from .logging_utils import get_logger
from .parser import Article
from .retrieval import ArticleIndex

log = get_logger(__name__)


# ── shared JSON parsing (tolerant of code fences) ──────────────────────────

def _parse_json_block(raw: str) -> dict:
    """Parse a JSON object out of raw model output, tolerant of ```fences."""
    s = raw.strip()
    if s.startswith("```"):
        s = s.split("```", 2)[1]
        if s.lstrip().lower().startswith("json"):
            s = s.split("\n", 1)[1] if "\n" in s else ""
        s = s.rsplit("```", 1)[0]
    start = s.find("{")
    end = s.rfind("}")
    if start == -1 or end == -1:
        raise ValueError(f"no JSON object in output: {raw[:200]}")
    return json.loads(s[start : end + 1])


def _format_articles_compact(pairs: list[tuple[Article, float]]) -> str:
    """Short article block for Opus — citation + first 300 chars of body.

    The full analytical prompt in ``brain.py`` uses the richer formatter;
    the pro features work on narrower inputs so we keep the prompt tight.
    """
    if not pairs:
        return "(asnjë nen i gjetur)"
    chunks: list[str] = []
    for a, _ in pairs:
        body = (a.body or "").strip().replace("\n", " ")
        if len(body) > 300:
            body = body[:297] + "..."
        chunks.append(f"• {a.citation} — {a.heading}\n  {body}")
    return "\n".join(chunks)


# ══════════════════════════════════════════════════════════════════════════
#  ① STRESS-TEST UDIENZA (red-team pre-udienza)
# ══════════════════════════════════════════════════════════════════════════

STRESS_TEST_SYSTEM = textwrap.dedent("""\
    Je avokat shqiptar senior i palës kundërshtare. Detyra: stres-teston
    rastin që tregohet poshtë, duke vepruar si "djalli i avokatisë"
    (devil's advocate). Qëllimi nuk është të fitosh argumentin, por t'i
    japësh avokatit që e ndjek rastin listën e plotë të dobësive që duhet
    të mbulojë para seancës.

    Ton profesional, shqipe juridike. Cito vetëm nene të Kodeve shqiptare
    që të janë vënë në dispozicion. Mos shpik jurisprudencë.

    Dorëzo VETËM një bllok JSON me këtë skemë:
    {
      "counter_brief": "string (250-400 fjalë, argumenti i kundërshtarit)",
      "weaknesses": [
        {"type": "factual|legal|procedural",
         "point": "string", "severity": "low|medium|high",
         "why_it_hurts": "string"}
      ],
      "cross_examination": [
        {"q": "string", "target": "dëshmitar|ekspert|klient|palë",
         "trap": "string (çfarë e bën pyetjen kurth)"}
      ],
      "procedural_objections": [
        {"objection": "string", "article": "string (Neni X i Kodit Y)",
         "timing": "paraprake|meritore"}
      ],
      "adverse_jurisprudence": [
        {"cite": "string (emër vendimi, vit)",
         "how_it_hurts": "string (pse e pret rrugën)"}
      ],
      "judges_questions": [
        "string (pyetje që gjykatësi i kujdesshëm do bëjë)"
      ],
      "score": {
        "winnability": 0-100,
        "risk_level": "low|medium|high",
        "verdict_summary": "string (një fjali, vlerësimi i përgjithshëm)"
      }
    }

    Rregulla: minimum 3 weaknesses, 8-12 cross_examination, 2-4 objections,
    1-3 adverse_jurisprudence, 3-5 judges_questions. Pa preambul, pa
    shpjegime jashtë JSON.
""")


def stress_test_hearing(
    backend: LLMBackend,
    index: ArticleIndex,
    hypothesis: str,
    *,
    case_docs: list[dict] | None = None,
) -> dict:
    """Run a single-pass red-team simulation on ``hypothesis``.

    Retrieves the top ~12 relevant articles via BM25 so the opposing
    counsel persona cites real law, then asks Opus for the structured
    JSON red-team. Returns the parsed dict; raises on malformed output.
    """
    retrieved = index.search(hypothesis, top_k=12)
    articles_block = _format_articles_compact(retrieved)
    docs_block = ""
    if case_docs:
        names = "\n".join(f"  • {d.get('filename', '?')}"
                          for d in case_docs)
        docs_block = f"\nDOKUMENTET E DOSJES:\n{names}\n"
    prompt = textwrap.dedent(f"""\
        Rasti i avokatit:
        \"\"\"{hypothesis}\"\"\"
        {docs_block}
        Nene të kodeve shqiptare që mund të jenë relevante:
        {articles_block}

        Gjeneroj JSON-in e stres-testit sipas skemës së sistemit.
    """)
    attachments: list[Path] = []
    for d in (case_docs or []):
        sp = d.get("storage_path")
        if sp and Path(sp).exists():
            attachments.append(Path(sp))

    raw = backend.complete(
        system=STRESS_TEST_SYSTEM,
        messages=[{"role": "user", "content": prompt}],
        max_tokens=4000,
        fast=False,
        session_id=None,
        attachments=attachments or None,
    )
    data = _parse_json_block(raw)
    # Light normalisation — keep downstream renderers forgiving.
    data.setdefault("weaknesses", [])
    data.setdefault("cross_examination", [])
    data.setdefault("procedural_objections", [])
    data.setdefault("adverse_jurisprudence", [])
    data.setdefault("judges_questions", [])
    data.setdefault("score", {})
    data["retrieved_articles"] = [
        {"citation": a.citation, "code": a.code, "number": a.number,
         "heading": a.heading, "score": round(s, 2)}
        for a, s in retrieved[:8]
    ]
    return data


# ══════════════════════════════════════════════════════════════════════════
#  ② AUDITOR DELLE CITAZIONI
# ══════════════════════════════════════════════════════════════════════════

# "Neni 130 i Kodit Civil" / "neni 83/a KPP" / "neni 1 i Kushtetutës"
# Captures the article number and the enclosing code name — greedy enough
# to grab multi-word titles like "Kodit të Procedurës Civile" but stops at
# sentence punctuation so it doesn't eat the rest of the sentence.
_STOP_WORDS = r"(?:dhe|edhe|ose|apo|ndërsa|si|ndaj|por|pra|ku|kur|që|për|nga|me|pa|nenit|neni|nenin|Neni|Nenit|Nenin|\d)"
_CITATION_RE = re.compile(
    r"[Nn]en(?:i|it|in|eve|e)?\s+"
    r"(?P<num>\d+(?:[/-][A-Za-z0-9çë]+)?)"
    r"\s+(?:i|e|të|t[eë])?\s*"
    r"(?P<code>"
        r"Kushtetut[ëaës]*"
        r"|Kod(?:i|it)"
            r"(?:\s+(?:t[eë]|i|e|së|s))?"
            r"(?:\s+(?!" + _STOP_WORDS + r"\b)[A-ZÇËa-zçëï]+){1,6}"
        r"|KPC|KPP|KC|KP|KF|KPA|KRR"
        r"|Ligj(?:i|it)\s+nr\.?\s*\d+[/-]\d+"
    r")",
    re.UNICODE,
)

# Map loose references ("KPC", "Kodi Civil") to config code ids. The full
# set is derived at runtime from LEGAL_DOCUMENTS so the matcher stays
# in sync with the corpus.

def _build_code_lookup() -> dict[str, str]:
    """Map every reasonable spelling of a code name → machine code id."""
    from .config import LEGAL_DOCUMENTS
    out: dict[str, str] = {}

    def _norm(s: str) -> str:
        return re.sub(r"\s+", " ", s.strip().lower())

    for d in LEGAL_DOCUMENTS:
        out[_norm(d.title_sq)] = d.code
        out[_norm(d.title_en)] = d.code
        out[_norm(d.code.replace("_", " "))] = d.code
        # common shorthand ("Kodi Civil", "Kodit Civil", "KPC", etc.)
        short = d.title_sq.replace("i Republikës së Shqipërisë", "").strip()
        out[_norm(short)] = d.code
        out[_norm(short.replace("Kodi ", "Kodit "))] = d.code
    # hard-coded abbreviations → codes
    abbr = {
        "kpc": "kodi_proc_civile",
        "kpp": "kodi_proc_penale",
        "kc": "kodi_civil",
        "kp": "kodi_penal",
        "kf": "kodi_familjes",
        "kpa": "kodi_proc_admin",
        "kushtetuta": "kushtetuta",
        "kushtetutës": "kushtetuta",
    }
    out.update(abbr)
    return out


_GENITIVE_WORDS = {"kodit", "kodi", "ligjit", "ligji", "kushtetutës",
                   "kushtetuta", "kushtetutes"}


def _normalize_code_key(s: str) -> str:
    """Normalize a code fragment: lowercase, strip punctuation, unify
    genitive forms (Kodit → Kodi), drop connector particles (të/i/e/së)
    and the boilerplate 'i Republikës së Shqipërisë' so synonyms collide."""
    s = re.sub(r"\s+", " ", s.strip().lower())
    s = s.rstrip(".,;:")
    s = re.sub(r"\bi\s+republik[eëa]s?\s+s[eë]\s+shqip[eë]ris[eë]\b", "", s)
    s = re.sub(r"\b(të|t[eë]|së|s[eë]|i|e)\b", " ", s)
    s = s.replace("kodit", "kodi").replace("ligjit", "ligji")
    s = s.replace("kushtetutës", "kushtetuta").replace("kushtetutes", "kushtetuta")
    s = re.sub(r"\s+", " ", s).strip()
    return s


def _resolve_code(fragment: str) -> str | None:
    """Return the machine code id for a free-text code fragment."""
    lookup = _build_code_lookup()
    # Add a normalized mirror of every key so the matcher is robust to
    # genitive / connector drift.
    norm_lookup: dict[str, str] = {
        _normalize_code_key(k): v for k, v in lookup.items()
    }
    key = _normalize_code_key(fragment)
    if key in norm_lookup:
        return norm_lookup[key]
    # prefix fallback — handle truncations both ways.
    best: tuple[int, str | None] = (0, None)
    for k, v in norm_lookup.items():
        if not k:
            continue
        if key.startswith(k) or k.startswith(key):
            if len(k) > best[0]:
                best = (len(k), v)
    return best[1]


def _find_article(
    index: ArticleIndex, code: str, number: str,
) -> Article | None:
    """Exact lookup of a specific (code, article number) in the corpus."""
    number = number.strip()
    for a in index.articles:
        if a.code == code and a.number == number:
            return a
    return None


AUDIT_SYSTEM = textwrap.dedent("""\
    Je redaktor juridik i rreptë. Për çdo citim që të jepet, detyra yte
    është të vendosësh nëse ai citim:
      - e mbështet me të vërtetë pohimin e autorit (status=correct);
      - ekziston por përdoret gabim (status=misapplied);
      - është shfuqizuar/zëvendësuar (status=superseded);
      - nuk ekziston me atë numër (status=not_found);
      - është i paqartë ose i duhet më shumë kontekst (status=unclear).

    Dorëzo VETËM një JSON:
    {
      "findings": [
        {"citation_text": "string (si ishte shkruar)",
         "resolved_code": "string | null",
         "resolved_number": "string | null",
         "status": "correct|misapplied|superseded|not_found|unclear",
         "verdict": "string (një fjali, pse)",
         "correct_version": "string | null",
         "note": "string | null"}
      ],
      "summary": {
        "total": int,
        "correct": int,
        "problematic": int,
        "overall_verdict": "string"
      }
    }

    Pa preambul, pa gjë jashtë JSON.
""")


def audit_citations(
    backend: LLMBackend,
    index: ArticleIndex,
    source_text: str,
) -> dict:
    """Scan ``source_text`` for legal citations, verify each against the
    corpus, then ask Opus to judge whether the citation actually supports
    the surrounding claim.
    """
    matches = list(_CITATION_RE.finditer(source_text))
    resolved: list[dict] = []
    for m in matches:
        number = m.group(1)
        code_frag = m.group(2)
        code_id = _resolve_code(code_frag)
        article = _find_article(index, code_id, number) if code_id else None
        # Pull a ±120-char window of surrounding context so the LLM sees
        # what claim the citation is attached to.
        start = max(0, m.start() - 180)
        end = min(len(source_text), m.end() + 180)
        context = source_text[start:end].strip()
        resolved.append({
            "citation_text": m.group(0),
            "resolved_code": code_id,
            "resolved_number": number,
            "found_in_corpus": article is not None,
            "corpus_heading": article.heading if article else None,
            "corpus_body": (article.body[:500] if article
                            else None),
            "repealed": bool(article.repealed) if article else None,
            "volatility": getattr(article, "volatility", None) if article
                           else None,
            "context": context,
        })

    if not resolved:
        return {
            "findings": [],
            "summary": {
                "total": 0, "correct": 0, "problematic": 0,
                "overall_verdict": "Asnjë citim i zbuluar në tekst.",
            },
        }

    prompt_body = json.dumps(resolved, ensure_ascii=False, indent=2)
    prompt = textwrap.dedent(f"""\
        Teksti burim (për referim):
        \"\"\"{source_text[:3000]}\"\"\"

        Citimet e zbuluara (me kontekst dhe me trupin e nenit real):
        {prompt_body}

        Vlerësoji të gjitha sipas skemës së JSON-it.
    """)
    raw = backend.complete(
        system=AUDIT_SYSTEM,
        messages=[{"role": "user", "content": prompt}],
        max_tokens=3500,
        fast=False,
        session_id=None,
    )
    data = _parse_json_block(raw)
    data.setdefault("findings", [])
    data.setdefault("summary", {})
    return data


# ══════════════════════════════════════════════════════════════════════════
#  ③ FABBRICA DEGLI ATTI PROCESSUALI
# ══════════════════════════════════════════════════════════════════════════

ACT_TYPES: dict[str, str] = {
    "padi":    "Padi (kërkesë-padi civile)",
    "kerkese": "Kërkesë drejtuar gjykatës",
    "ankim":   "Ankim (apel) kundër vendimit",
    "rekurs":  "Rekurs në Gjykatën e Lartë",
    "memorie": "Memorie sqaruese / konkluzione",
    "kerkese_penale": "Kërkesë në procedim penal",
    "kundershtim": "Kundërshtim (opozitë) ndaj ekzekutimit",
}


DRAFT_SYSTEM = textwrap.dedent("""\
    Je avokat shqiptar me përvojë që harton akte procedurale. Detyrë:
    nga përshkrimi që të jepet dhe nenet që të ofrohen, shkruaj një akt
    procedural SHQIPËRI, gati për depozitim në gjykatë.

    Strukturë standarde:
      1. Gjykata që i drejtohet
      2. Palët (paditësi / i padituri ose analogët)
      3. Lënda e kërkesës (një rresht)
      4. Rrethanat faktike (paragrafë të numëruar)
      5. Analiza juridike (cito nenet e dhëna me formatin "Neni X i
         Kodit Y"; mos shpik nene)
      6. KËRKOJMË (petitum) — lista e qartë e kërkesave
      7. Lista e provave
      8. Vendi, data, nënshkrimi

    Dorëzo VETËM një JSON:
    {
      "title": "string (titulli i aktit, p.sh. 'KËRKESË-PADI')",
      "court": "string",
      "parties": {
        "actor": "string (emërtimi i palës kërkuese)",
        "defendant": "string (emërtimi i palës tjetër, nëse ka)"
      },
      "subject_matter": "string",
      "body_markdown": "string (TEKSTI I PLOTË I AKTIT në shqip, gati
        për printim; përdor headings, lista, paragrafë të numëruar)",
      "petitum": ["string", "string", ...],
      "cited_articles": [
        {"code": "string", "number": "string",
         "citation": "Neni X i Kodit Y"}
      ],
      "warnings": [
        "string (çdo gjë që avokati duhet të kontrollojë para depozitimit:
         afate, taksë gjykate, dokumente shoqëruese që i duhen)"
      ]
    }

    Mos shto asgjë jashtë JSON-it.
""")


def draft_act(
    backend: LLMBackend,
    index: ArticleIndex,
    *,
    act_type: str,
    brief: str,
    case_docs: list[dict] | None = None,
) -> dict:
    """Generate a filing-ready procedural act from a free-text brief."""
    if act_type not in ACT_TYPES:
        raise ValueError(f"unknown act_type: {act_type!r}")
    retrieved = index.search(brief, top_k=15)
    articles_block = _format_articles_compact(retrieved)
    docs_block = ""
    attachments: list[Path] = []
    if case_docs:
        names = "\n".join(f"  • {d.get('filename', '?')}"
                          for d in case_docs)
        docs_block = f"\nDOKUMENTET E DOSJES (lexoji):\n{names}\n"
        for d in case_docs:
            sp = d.get("storage_path")
            if sp and Path(sp).exists():
                attachments.append(Path(sp))
    act_label = ACT_TYPES[act_type]
    prompt = textwrap.dedent(f"""\
        Lloji i aktit që duhet të hartohet: {act_label}

        Përshkrimi i rastit nga avokati:
        \"\"\"{brief}\"\"\"
        {docs_block}
        Nene relevante nga kodet shqiptare (përdor vetëm këto):
        {articles_block}

        Harto aktin duke respektuar strukturën e skemës JSON.
    """)
    raw = backend.complete(
        system=DRAFT_SYSTEM,
        messages=[{"role": "user", "content": prompt}],
        max_tokens=6000,
        fast=False,
        session_id=None,
        attachments=attachments or None,
    )
    data = _parse_json_block(raw)
    data.setdefault("petitum", [])
    data.setdefault("cited_articles", [])
    data.setdefault("warnings", [])
    data["act_type"] = act_type
    return data


def render_act_docx(draft: dict, out_path: Path) -> Path:
    """Serialise a drafted act to a .docx file using python-docx."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    # Title
    title = doc.add_heading(draft.get("title") or "AKT PROCEDURAL", level=0)
    title.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER is int 1
    court = draft.get("court") or ""
    if court:
        p = doc.add_paragraph()
        r = p.add_run(f"DREJTUAR: {court}")
        r.bold = True
    parties = draft.get("parties") or {}
    if parties.get("actor"):
        doc.add_paragraph(f"PALA KËRKUESE: {parties['actor']}")
    if parties.get("defendant"):
        doc.add_paragraph(f"PALA E PADITUR: {parties['defendant']}")
    subject = draft.get("subject_matter")
    if subject:
        p = doc.add_paragraph()
        r = p.add_run(f"LËNDA: {subject}")
        r.bold = True

    doc.add_paragraph()
    body = draft.get("body_markdown") or ""
    for line in body.split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith(("- ", "• ")):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        else:
            p = doc.add_paragraph(stripped)
            for run in p.runs:
                run.font.size = Pt(11)

    petitum = draft.get("petitum") or []
    if petitum:
        doc.add_heading("KËRKOJMË", level=1)
        for item in petitum:
            doc.add_paragraph(str(item), style="List Number")

    cited = draft.get("cited_articles") or []
    if cited:
        doc.add_heading("Bazë ligjore e cituar", level=2)
        for c in cited:
            doc.add_paragraph(
                f"• {c.get('citation') or c.get('code') + ' nr. ' + c.get('number', '')}"
            )

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


# ══════════════════════════════════════════════════════════════════════════
#  ④ CASCATA TERMINI PROCESSUALI (pure rules engine)
# ══════════════════════════════════════════════════════════════════════════
#
# Each "procedural event" (notification of sentence, service of summons,
# filing of appeal, etc.) triggers a cascade of downstream statutory
# deadlines. We encode the mapping deterministically here — no LLM call —
# so the output is auditable and reproducible. The LLM is optional for
# suggesting the proper event type from a free-text description.

# ISO-tagged rulebook. Each rule: (key, label_sq, code, base_article,
# days_count, counting_method, urgency, notes_sq).
#
# counting_method:
#   "calendar" — every day counts
#   "business" — Mon-Fri only (approximation: no holiday calendar here)
#
# base_article is the primary article in KPC/KPP that grounds the rule.
# Keep days conservative — if the lawyer needs certainty they verify the
# article directly from the linked panel.

DEADLINE_RULES: dict[str, list[dict]] = {
    # Civil — sentenza di primo grado notificata
    "njoftim_vendimi_civil_shkalle_pare": [
        {"key": "ankim_civil", "label": "Afati për ankim (apel)",
         "code": "kodi_proc_civile", "article": "443",
         "days": 15, "counting": "calendar",
         "urgency": "high",
         "notes": "Ankimi paraqitet në gjykatën që ka dhënë vendimin. "
                  "Tejkalimi i afatit e bën vendimin të formës së prerë."},
        {"key": "kerkese_ekzekutimi", "label": "Mund të kërkohet ekzekutimi",
         "code": "kodi_proc_civile", "article": "510",
         "days": 15, "counting": "calendar",
         "urgency": "medium",
         "notes": "Pas kalimit të afatit të ankimit, vendimi bëhet titull "
                  "ekzekutiv (po qe se ligji e lejon ekzekutimin para "
                  "formës së prerë, shih nenet përkatëse)."},
    ],
    # Civil — sentenza d'appello notificata
    "njoftim_vendimi_apeli": [
        {"key": "rekurs_gjykata_larte", "label": "Afati për rekurs në Gjykatën e Lartë",
         "code": "kodi_proc_civile", "article": "472",
         "days": 30, "counting": "calendar",
         "urgency": "high",
         "notes": "Rekursi depozitohet në Gjykatën e Apelit që ka dhënë "
                  "vendimin. Duhet baza ligjore për çdo shkak rekursi."},
    ],
    # Civil — sentenza definitiva (forma e prerë)
    "njoftim_vendimi_prere_civil": [
        {"key": "rishikim_neni_494", "label": "Afati për rishikim",
         "code": "kodi_proc_civile", "article": "494",
         "days": 30, "counting": "calendar",
         "urgency": "medium",
         "notes": "Rishikimi kërkohet brenda 30 ditëve nga data që pala "
                  "ka marrë dijeni për shkaqet e rishikimit."},
    ],
    # Civil — notifica di un atto di esecuzione
    "njoftim_urdher_ekzekutimi": [
        {"key": "kundershtim_ekzekutimi", "label": "Kundërshtim ndaj ekzekutimit",
         "code": "kodi_proc_civile", "article": "610",
         "days": 5, "counting": "calendar",
         "urgency": "critical",
         "notes": "Afat jashtëzakonisht i shkurtër: paraqitet në gjykatën e "
                  "vendit të ekzekutimit. Humbja = rrezikon ekzekutimin."},
    ],
    # Penale — sentenza di primo grado (KPP)
    "njoftim_vendimi_penal_shkalle_pare": [
        {"key": "ankim_penal", "label": "Afati për ankim (KPP)",
         "code": "kodi_proc_penale", "article": "410",
         "days": 15, "counting": "calendar",
         "urgency": "high",
         "notes": "Ankimi depozitohet në gjykatën që ka dhënë vendimin. "
                  "Afati fillon nga dita e njoftimit ose e shpalljes."},
    ],
    # Penale — sentenza d'appello penale
    "njoftim_vendimi_apel_penal": [
        {"key": "rekurs_penal_gjykata_larte", "label": "Afati për rekurs në Gjykatën e Lartë (penale)",
         "code": "kodi_proc_penale", "article": "432",
         "days": 30, "counting": "calendar",
         "urgency": "high",
         "notes": "Shkaqet e rekursit janë të kufizuara me ligj. Verifiko "
                  "nenin 432 të KPP-së."},
    ],
    # Penale — masa parasegura / arrestimi
    "zbatim_mase_sigurie": [
        {"key": "ankim_mase_sigurie", "label": "Ankim ndaj masës së sigurimit",
         "code": "kodi_proc_penale", "article": "249",
         "days": 5, "counting": "calendar",
         "urgency": "critical",
         "notes": "Afat shumë i shkurtër — prek lirinë personale."},
    ],
    # Civile — thirrje në gjyq (prima udienza)
    "thirrje_seance": [
        {"key": "parashtrim_mbrojtjes", "label": "Paraqitja e mbrojtjes",
         "code": "kodi_proc_civile", "article": "163",
         "days": 0, "counting": "calendar",
         "urgency": "high",
         "notes": "Mbrojtja paraqitet deri në ditën e parë të gjykimit; "
                  "aq sa më herët, aq më mirë — afati tregohet vetëm "
                  "me datën e seancës."},
    ],
    # Amministrativo — atto impugnabile
    "njoftim_akti_administrativ": [
        {"key": "ankim_administrativ", "label": "Ankim administrativ",
         "code": "kodi_proc_admin", "article": "135",
         "days": 30, "counting": "calendar",
         "urgency": "medium",
         "notes": "Ankimi drejtuar organit epror ose Gjykatës "
                  "Administrative të Shkallës së Parë, sipas llojit të aktit."},
    ],
}


EVENT_TYPE_LABELS_SQ: dict[str, str] = {
    "njoftim_vendimi_civil_shkalle_pare": "Njoftim i vendimit civil — shkalla e parë",
    "njoftim_vendimi_apeli": "Njoftim i vendimit të apelit (civil)",
    "njoftim_vendimi_prere_civil": "Vendim civil i formës së prerë",
    "njoftim_urdher_ekzekutimi": "Njoftim i urdhrit të ekzekutimit",
    "njoftim_vendimi_penal_shkalle_pare": "Njoftim i vendimit penal — shkalla e parë",
    "njoftim_vendimi_apel_penal": "Njoftim i vendimit të apelit (penal)",
    "zbatim_mase_sigurie": "Zbatim i masës së sigurimit",
    "thirrje_seance": "Thirrje në seancë gjyqësore",
    "njoftim_akti_administrativ": "Njoftim i aktit administrativ",
}


def _add_days(base: date, days: int, counting: str) -> date:
    if counting == "business":
        out = base
        added = 0
        while added < days:
            out = out + timedelta(days=1)
            if out.weekday() < 5:
                added += 1
        return out
    return base + timedelta(days=days)


def compute_deadline_cascade(
    event_type: str, event_date: str,
) -> dict:
    """Given a procedural event type + ISO date, return the cascade.

    Pure rules engine: no LLM, no I/O. Output is stable and auditable.
    """
    if event_type not in DEADLINE_RULES:
        raise ValueError(f"unknown event_type: {event_type!r}")
    try:
        base = date.fromisoformat(event_date)
    except ValueError as exc:
        raise ValueError(f"event_date must be ISO YYYY-MM-DD: {exc}")

    today = date.today()
    derived: list[dict] = []
    for rule in DEADLINE_RULES[event_type]:
        due = _add_days(base, rule["days"], rule["counting"])
        days_left = (due - today).days
        if days_left < 0:
            status = "expired"
        elif days_left == 0:
            status = "due_today"
        elif days_left <= 2:
            status = "imminent"
        elif days_left <= 7:
            status = "soon"
        else:
            status = "upcoming"
        derived.append({
            "key": rule["key"],
            "label": rule["label"],
            "code": rule["code"],
            "article": rule["article"],
            "citation": f"Neni {rule['article']} i "
                        f"{_code_title_sq(rule['code'])}",
            "days": rule["days"],
            "counting": rule["counting"],
            "base_date": event_date,
            "due_date": due.isoformat(),
            "days_left": days_left,
            "status": status,
            "urgency": rule["urgency"],
            "notes": rule.get("notes", ""),
        })
    derived.sort(key=lambda d: d["due_date"])
    return {
        "event_type": event_type,
        "event_label": EVENT_TYPE_LABELS_SQ.get(event_type, event_type),
        "event_date": event_date,
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "derived_deadlines": derived,
    }


def _code_title_sq(code: str) -> str:
    """Human Albanian title for a code id (for citations in output)."""
    from .config import LEGAL_DOCUMENTS
    for d in LEGAL_DOCUMENTS:
        if d.code == code:
            return d.title_sq
    return code


def cascade_event_types() -> list[dict]:
    """Return the list of supported cascade event types for the UI."""
    out: list[dict] = []
    for key, label in EVENT_TYPE_LABELS_SQ.items():
        rules = DEADLINE_RULES[key]
        out.append({
            "key": key,
            "label": label,
            "rule_count": len(rules),
            "rules_preview": [
                {"key": r["key"], "days": r["days"], "urgency": r["urgency"]}
                for r in rules
            ],
        })
    return out


# ══════════════════════════════════════════════════════════════════════════
#  ⑤ TIMELINE DEL FASCICOLO (V7.12 feature A)
# ══════════════════════════════════════════════════════════════════════════
#
# Given a case description + every document attached to the dossier, build
# a chronological reconstruction of every dated event, flag contradictions
# between sources, and surface temporal gaps. The output is a single JSON
# blob the UI renders as a vertical timeline with click-to-source links.

TIMELINE_SYSTEM = textwrap.dedent("""\
    Je një jurist analist që ndërton kronologjinë e plotë të një rasti
    ligjor shqiptar duke nxjerrë çdo ngjarje të datuar nga dokumentet e
    dosjes dhe nga përshkrimi i avokatit.

    Detyra:
    1) Lexo me kujdes përshkrimin e rastit dhe dokumentet bashkëngjitur.
    2) Nxirr ÇDO ngjarje me datë: nënshkrim kontrate, njoftime, pushime,
       pagesa, mesazhe, takime, seanca, dëmtime, vërejtje, etj.
    3) Identifiko KONTRADIKTAT — kur dy burime japin data të ndryshme për
       të njëjtën ngjarje, ose kur faktet bien në kundërshtim.
    4) Identifiko BOSHLLËQET — periudha të gjata pa veprim ku ligji do të
       priste një reagim (p.sh. 30 ditë heshtje pas një njoftimi formal).
    5) Cito vetëm dokumentet që janë dhënë; mos shpik prova.

    Përgjigja DUHET të jetë vetëm një objekt JSON i vlefshëm, pa tekst
    shtesë jashtë tij, që ndjek këtë skemë:

    {
      "events": [
        {
          "date": "YYYY-MM-DD",
          "date_confidence": "exact" | "approximate" | "inferred",
          "time": "HH:MM" | null,
          "type": "njoftim" | "kontrate" | "pushim" | "pagese" |
                  "mesazh" | "seance" | "demti" | "vendim" | "tjeter",
          "summary": "përshkrim i shkurtër në shqip (1-2 fjali)",
          "parties": ["pala 1", "pala 2"],
          "source_doc": "filename ose 'description' nëse vjen nga teksti",
          "source_excerpt": "citim i drejtpërdrejtë nga burimi (max 200 char)",
          "legal_significance": "pse kjo ngjarje ka rëndësi juridike (1 fjali)"
        }
      ],
      "contradictions": [
        {
          "issue": "çfarë nuk përputhet (p.sh. data e pushimit)",
          "claims": [
            {"value": "vlera e pretenduar", "source": "dokument-X.pdf"},
            {"value": "vlera tjetër", "source": "dokument-Y.pdf"}
          ],
          "severity": "high" | "medium" | "low",
          "tactical_note": "si mund ta përdorim ose neutralizojmë"
        }
      ],
      "gaps": [
        {
          "from": "YYYY-MM-DD",
          "to": "YYYY-MM-DD",
          "duration_days": 0,
          "concern": "pse ky boshllëk është i dyshimtë juridikisht"
        }
      ],
      "summary": "një paragraf përmbledhës i kronologjisë në shqip"
    }

    Renditi events sipas datës rritëse. Datat duhet të jenë në format
    ISO (YYYY-MM-DD). Nëse një datë është e përafërt (p.sh. "në mars"),
    përdor ditën e parë të muajit dhe shëno date_confidence: "approximate".
    Nëse asnjë ngjarje nuk gjendet, kthe events: [] me një summary që e
    shpjegon.
""")


def build_case_timeline(
    backend: LLMBackend,
    case_summary: str,
    case_title: str,
    case_docs: list[dict] | None = None,
) -> dict:
    """Reconstruct the chronological timeline of a case.

    ``case_summary`` is typically the first user message of the case (the
    lawyer's narrative); ``case_docs`` are the dossier attachments. We send
    every dossier file to the backend so Opus can read them directly.
    """
    docs_block = ""
    attachments: list[Path] = []
    if case_docs:
        names: list[str] = []
        for d in case_docs:
            fn = d.get("filename") or "?"
            names.append(f"  • {fn}")
            sp = d.get("storage_path")
            if sp and Path(sp).exists():
                attachments.append(Path(sp))
        docs_block = f"\nDOKUMENTET E DOSJES (bashkëngjitur):\n" + "\n".join(names) + "\n"

    prompt = textwrap.dedent(f"""\
        EMRI I RASTIT: {case_title}

        PËRSHKRIMI / NARRATIVA E AVOKATIT:
        \"\"\"{case_summary}\"\"\"
        {docs_block}
        Ndërto kronologjinë sipas skemës JSON të sistemit. Përgjigja vetëm
        si objekt JSON, asgjë tjetër.
    """)

    raw = backend.complete(
        system=TIMELINE_SYSTEM,
        messages=[{"role": "user", "content": prompt}],
        max_tokens=6000,
        fast=False,
        session_id=None,
        attachments=attachments or None,
    )
    data = _parse_json_block(raw)

    # Normalise + sort.
    events = data.get("events") or []
    for ev in events:
        ev.setdefault("date_confidence", "exact")
        ev.setdefault("parties", [])
        ev.setdefault("source_doc", "description")
        ev.setdefault("source_excerpt", "")
        ev.setdefault("legal_significance", "")
    # Stable chronological order; events without parseable date fall last.
    def _sort_key(e: dict) -> tuple[int, str, str]:
        d = (e.get("date") or "").strip()
        try:
            datetime.strptime(d, "%Y-%m-%d")
            return (0, d, e.get("time") or "")
        except ValueError:
            return (1, d, "")
    events.sort(key=_sort_key)
    data["events"] = events
    data.setdefault("contradictions", [])
    data.setdefault("gaps", [])
    data.setdefault("summary", "")
    data["generated_at"] = datetime.now(timezone.utc).isoformat(timespec="seconds")
    data["doc_count"] = len(case_docs or [])
    data["event_count"] = len(events)
    return data


# ══════════════════════════════════════════════════════════════════════════
#  ⑥ ADVERSARIAL LOOP (V7.12 feature C)
# ══════════════════════════════════════════════════════════════════════════
#
# Iterative red-team: round-by-round, the AI argues both sides until no new
# substantive attack can be raised. Distinct from the single-pass stress
# test in that the LAWYER's strategy evolves between rounds — every defence
# becomes input to the next attack. Output is a sequence of rounds plus a
# final converged plan.

ADVERSARIAL_ATTACK_SYSTEM = textwrap.dedent("""\
    Je avokat shqiptar senior i palës KUNDËRSHTARE. Ke marrë strategjinë
    e fundit të avokatit përballë + raundet e mëparshme. Detyra: gjej
    sulmin më të fortë që NUK është adresuar ende dhe që do të shqetësonte
    avokatin më shumë.

    Përgjigja DUHET të jetë vetëm një objekt JSON, sipas kësaj skeme:

    {
      "round": <numri i raundit>,
      "attack_type": "ligjor" | "procedural" | "faktik" | "evidencë" | "jurisprudencë",
      "attack_thesis": "një fjali që përshkruan sulmin",
      "attack_argumentation": "argumentimi i plotë (3-5 paragrafë në shqip juridik)",
      "cited_articles": ["Neni X i KP", ...],
      "risk_to_lawyer": "high" | "medium" | "low",
      "novelty": "i ri / shtjellim i një sulmi më të hershëm",
      "if_unanswered": "çfarë ndodh në sallë nëse avokati nuk e kundërpërgjigj"
    }

    Mos përsërit sulme nga raundet e mëparshme. Nëse vërtet nuk ka më
    sulm domethënës, vendos attack_thesis: "konvergjuar" dhe risk_to_lawyer:
    "low" — kjo i thotë sistemit të ndalojë.
""")

ADVERSARIAL_DEFENSE_SYSTEM = textwrap.dedent("""\
    Je avokat shqiptar senior që mbron klientin. Sapo ke pranuar sulmin
    më të fundit nga pala kundërshtare. Detyra: jep kundërpërgjigjen më
    të fortë të mundshme, duke pranuar atë që duhet pranuar dhe duke e
    konvertuar dobësinë në avantazh kur ka mundësi.

    Përgjigja DUHET të jetë vetëm një objekt JSON, sipas skemës:

    {
      "round": <numri i raundit>,
      "defense_thesis": "një fjali që përshkruan kundërpërgjigjen",
      "defense_argumentation": "argumentimi i plotë (3-5 paragrafë)",
      "cited_articles": ["Neni X i KP", ...],
      "concession": "çfarë pranojmë (nëse diçka)",
      "residual_risk": "high" | "medium" | "low",
      "remaining_weakness": "çfarë mbetet e pambrojtur edhe pas kësaj"
    }
""")

ADVERSARIAL_SUMMARY_SYSTEM = textwrap.dedent("""\
    Je një jurist supervizor që mbledh raundet e betejës dhe nxjerr
    strategjinë finale për avokatin. Përgjigja vetëm si JSON:

    {
      "final_strategy": "paragrafi i strategjisë së kalitur — çfarë të bëjë",
      "ranked_action_items": [
        {"priority": 1, "action": "...", "deadline_relative": "p.sh. para seancës"},
        ...
      ],
      "fortified_positions": ["pikat e tezës që rezistojnë sulmeve"],
      "remaining_vulnerabilities": ["dobësitë që nuk u shuan plotësisht"],
      "verdict_likelihood": "favorable" | "uncertain" | "unfavorable",
      "key_takeaway": "një fjali që e mban mend avokati hyrë në sallë"
    }
""")


def adversarial_loop(
    backend: LLMBackend,
    index: ArticleIndex,
    hypothesis: str,
    max_rounds: int = 5,
    *,
    case_docs: list[dict] | None = None,
) -> dict:
    """Run an iterative attacker/defender loop on ``hypothesis``.

    Each round = 1 attack + 1 defence (2 Opus calls). Stops early when the
    attacker outputs ``attack_thesis: "konvergjuar"`` or after ``max_rounds``.
    A final summary call distils everything into the strategist's plan.
    """
    retrieved = index.search(hypothesis, top_k=10)
    articles_block = _format_articles_compact(retrieved)
    docs_block = ""
    if case_docs:
        names = "\n".join(f"  • {d.get('filename', '?')}" for d in case_docs)
        docs_block = f"\nDOKUMENTET E DOSJES:\n{names}\n"
    attachments: list[Path] = []
    for d in (case_docs or []):
        sp = d.get("storage_path")
        if sp and Path(sp).exists():
            attachments.append(Path(sp))

    rounds: list[dict] = []
    history_lines: list[str] = []
    current_strategy = hypothesis

    for r in range(1, max_rounds + 1):
        history_block = ("\n".join(history_lines)
                         if history_lines else "(asnjë raund i mëparshëm)")
        # ── attacker turn ──
        attack_prompt = textwrap.dedent(f"""\
            Raundi #{r}.
            Strategjia e fundit e avokatit:
            \"\"\"{current_strategy}\"\"\"
            {docs_block}
            Nene relevante:
            {articles_block}

            HISTORIA E RAUNDEVE TË KALUARA:
            {history_block}

            Gjeneroj sulmin për këtë raund sipas skemës JSON.
        """)
        atk_raw = backend.complete(
            system=ADVERSARIAL_ATTACK_SYSTEM,
            messages=[{"role": "user", "content": attack_prompt}],
            max_tokens=2500,
            fast=False,
            session_id=None,
            attachments=attachments or None,
        )
        try:
            attack = _parse_json_block(atk_raw)
        except Exception:
            log.exception("adversarial: attack JSON parse failed at round %d", r)
            break
        attack["round"] = r
        thesis = (attack.get("attack_thesis") or "").strip().lower()
        if "konvergj" in thesis or attack.get("risk_to_lawyer") == "low" and r > 1:
            rounds.append({"round": r, "attack": attack, "defense": None,
                           "converged": True})
            log.info("adversarial: converged at round %d", r)
            break

        # ── defender turn ──
        defense_prompt = textwrap.dedent(f"""\
            Raundi #{r}. Sulmi i sapomarrë:
            \"\"\"{json.dumps(attack, ensure_ascii=False, indent=2)}\"\"\"

            Strategjia e fundit e avokatit (përpara sulmit):
            \"\"\"{current_strategy}\"\"\"

            Jep kundërpërgjigjen sipas skemës JSON.
        """)
        def_raw = backend.complete(
            system=ADVERSARIAL_DEFENSE_SYSTEM,
            messages=[{"role": "user", "content": defense_prompt}],
            max_tokens=2500,
            fast=False,
            session_id=None,
        )
        try:
            defense = _parse_json_block(def_raw)
        except Exception:
            log.exception("adversarial: defense JSON parse failed at round %d", r)
            rounds.append({"round": r, "attack": attack, "defense": None,
                           "converged": False})
            break
        defense["round"] = r
        rounds.append({"round": r, "attack": attack, "defense": defense,
                       "converged": False})

        # Update strategy = previous + this defence (compact summary).
        history_lines.append(
            f"Raundi {r}: SULMI = {attack.get('attack_thesis', '')[:200]} | "
            f"MBROJTJA = {defense.get('defense_thesis', '')[:200]}"
        )
        current_strategy = (
            f"{hypothesis}\n\n"
            f"[Pas raundit {r}] {defense.get('defense_thesis', '')}\n"
            f"{defense.get('defense_argumentation', '')[:500]}"
        )

    # ── final summary ──
    summary_prompt = textwrap.dedent(f"""\
        Hipoteza fillestare:
        \"\"\"{hypothesis}\"\"\"

        Raundet e betejës (sulm + mbrojtje):
        {json.dumps(rounds, ensure_ascii=False, indent=2)}

        Distilo strategjinë finale sipas skemës JSON.
    """)
    sum_raw = backend.complete(
        system=ADVERSARIAL_SUMMARY_SYSTEM,
        messages=[{"role": "user", "content": summary_prompt}],
        max_tokens=2500,
        fast=False,
        session_id=None,
    )
    try:
        summary = _parse_json_block(sum_raw)
    except Exception:
        log.exception("adversarial: summary JSON parse failed")
        summary = {
            "final_strategy": "(përmbledhja dështoi — shih raundet e plota)",
            "ranked_action_items": [],
            "fortified_positions": [],
            "remaining_vulnerabilities": [],
            "verdict_likelihood": "uncertain",
            "key_takeaway": "",
        }

    return {
        "hypothesis": hypothesis,
        "rounds": rounds,
        "round_count": len(rounds),
        "summary": summary,
        "retrieved_articles": [
            {"citation": a.citation, "code": a.code, "number": a.number,
             "heading": a.heading, "score": round(s, 2)}
            for a, s in retrieved[:8]
        ],
        "generated_at": datetime.now(timezone.utc).isoformat(timespec="seconds"),
    }


# ───────────────────────── V7.12 — Strategy compass ─────────────────────────
# Generate a branching decision tree for a case: at each fork, the lawyer can
# pick path A or B (e.g. "padit në gjykatë" vs "kërko ndërmjetësim"), with
# probability of success, estimated cost/duration, and reasoning per branch.
# The model also flags the recommended path and dead-ends.

STRATEGY_COMPASS_SYSTEM = textwrap.dedent("""\
    Je një konsulent strategjik juridik shqiptar. Detyra: nisur nga objektivi
    dhe situata e dosjes, ndërto një PEMË VENDIMARRJE me degë alternative.
    Çdo nyje paraqet një vendim taktik (p.sh. "Padit në gjykatë" vs "Kërko
    ndërmjetësim") me probabilitete dhe kosto të vlerësuara realisht në bazë
    të KP/KPC/KPP shqiptar.

    Përgjigja vetëm si JSON valid:

    {
      "objective": "objektivi final (p.sh. 'shpërblim dëmi 5M lekë')",
      "root": {
        "id": "n0",
        "label": "Situata fillestare",
        "description": "një paragraf që përshkruan kontekstin",
        "type": "start"
      },
      "branches": [
        {
          "id": "n1",
          "parent_id": "n0",
          "label": "Hapi i parë i mundshëm (folje + objekt)",
          "description": "shpjegim 2-3 fjali çfarë presupozon dega",
          "type": "decision" | "action" | "outcome",
          "probability_success": 0.0-1.0,
          "estimated_cost_alm": "p.sh. 50000-150000",
          "estimated_duration": "p.sh. 6-12 muaj",
          "legal_basis": ["KPC neni X", "..."],
          "pros": ["arsye favorizuese"],
          "cons": ["arsye kundër"],
          "next_ids": ["n2", "n3"],
          "is_recommended": true | false,
          "is_dead_end": false
        }
      ],
      "recommended_path": ["n0", "n1", "n3"],
      "key_insights": ["mësimi 1", "mësimi 2"],
      "warnings": ["risk që duhet shmangur"]
    }

    Rregulla:
    - Maksimumi 12 nyje totale (root + branches), 3 nivele thellësi.
    - Çdo degë vendimi (decision) MUND të ketë 2-3 fëmijë; degët veprim/result
      janë gjethe (next_ids = []).
    - Probability_success: 0.0 = pamundur, 1.0 = e sigurt; bazuar tek prova
      tipike dhe juridiksioni i krahasueshëm.
    - Recommended_path = sekuenca optimale e id-ve nga root tek gjethja
      më e mirë; recommended branches duhet të kenë "is_recommended": true.
    - dead_end = true për degët që duhen shmangur (kosto/risk i lartë).
""")


def build_strategy_compass(
    backend: LLMBackend,
    index: ArticleIndex,
    objective: str,
    case_summary: str,
    *,
    case_title: str | None = None,
    case_docs: list[dict] | None = None,
) -> dict:
    """Generate a decision-tree compass for ``objective`` given the case state.

    Returns a dict with the JSON shape declared in ``STRATEGY_COMPASS_SYSTEM``,
    plus ``meta`` (node_count, depth, generated_at, retrieved_articles).
    """
    retrieved = index.search(f"{objective}\n\n{case_summary}", top_k=10)
    articles_block = _format_articles_compact(retrieved)

    docs_block = ""
    if case_docs:
        parts = []
        for d in case_docs[:5]:
            name = d.get("filename") or d.get("name") or "dokument"
            excerpt = (d.get("text") or "")[:600]
            if excerpt:
                parts.append(f"### {name}\n{excerpt}")
        if parts:
            docs_block = "\n\nDokumentet e dosjes (ekstrakt):\n" + "\n\n".join(parts)

    title_line = f"Titulli i dosjes: {case_title}\n" if case_title else ""
    user_prompt = textwrap.dedent(f"""\
        {title_line}Objektivi i klientit: {objective}

        Përmbledhje e dosjes:
        {case_summary}
        {docs_block}

        Nenet relevante nga KP/KPC/KPP/Kushtetuta:
        {articles_block}

        Ndërto pemën e vendimarrjes. JSON only.
    """)

    raw = backend.complete(
        system=STRATEGY_COMPASS_SYSTEM,
        messages=[{"role": "user", "content": user_prompt}],
        max_tokens=4500,
    )
    data = _parse_json_block(raw)

    branches = data.get("branches") or []
    root = data.get("root") or {}
    node_count = len(branches) + (1 if root else 0)

    parent_to_depth: dict[str, int] = {}
    if root.get("id"):
        parent_to_depth[root["id"]] = 0
    changed = True
    while changed:
        changed = False
        for b in branches:
            pid = b.get("parent_id")
            bid = b.get("id")
            if pid in parent_to_depth and bid not in parent_to_depth:
                parent_to_depth[bid] = parent_to_depth[pid] + 1
                changed = True
    max_depth = max(parent_to_depth.values()) if parent_to_depth else 0

    data["meta"] = {
        "node_count": node_count,
        "depth": max_depth,
        "generated_at": datetime.now(timezone.utc).isoformat(timespec="seconds"),
        "retrieved_articles": [
            {"citation": a.citation, "code": a.code, "number": a.number,
             "heading": a.heading, "score": round(s, 2)}
            for a, s in retrieved[:8]
        ],
    }
    return data


# ── V8.11 Citation Shield V2 — provenance docx export ────────────────────

def provenance_docx(pack: dict) -> bytes:
    """Render a ProvenancePack as a .docx audit document.

    The output is suitable for filing in the case dossier: a one-page
    record of WHO produced WHAT WHEN against WHICH KB version, plus the
    full list of citations and retrieved articles so a reader can audit
    the answer years later — even after the live KB has rotated.
    """
    from docx import Document
    from docx.shared import Pt
    from io import BytesIO

    doc = Document()

    title = doc.add_heading("PROVENANCE PACK", level=0)
    title.alignment = 1

    pid = pack.get("response_id") or "—"
    ts = pack.get("timestamp_iso") or "—"
    juris = pack.get("jurisdiction") or "AL"
    p = doc.add_paragraph()
    r = p.add_run(f"ID: {pid}    ·    {ts}    ·    {juris}")
    r.font.size = Pt(9)

    if pack.get("refused"):
        warning = doc.add_paragraph()
        wr = warning.add_run(
            "⚠ REFUSAL — citimet nuk u verifikuan. "
            "Përgjigjja ruhet si e pasigurt."
        )
        wr.bold = True

    doc.add_heading("Konfigurim modeli", level=2)
    cfg_rows = [
        ("Modeli", pack.get("model") or "—"),
        ("Versioni i system prompt", pack.get("system_prompt_version") or "—"),
        ("Versioni i bazës (KB)", pack.get("kb_version") or "—"),
        ("Hash kërkese (input)", pack.get("prompt_hash") or "—"),
        ("Hash përgjigjeje (output)", pack.get("response_hash") or "—"),
        ("Besimi (confidence)",
         f"{pack.get('confidence', 0)} · {pack.get('confidence_label') or '—'}"),
    ]
    table = doc.add_table(rows=len(cfg_rows), cols=2)
    table.style = "Light Grid Accent 1"
    for i, (k, v) in enumerate(cfg_rows):
        cells = table.rows[i].cells
        cells[0].text = k
        cells[1].text = str(v)

    citations = pack.get("citations") or {}
    items = citations.get("items") or []
    if items:
        doc.add_heading("Citime dhe verifikim", level=2)
        cit_table = doc.add_table(rows=1, cols=3)
        cit_table.style = "Light Grid Accent 1"
        hdr = cit_table.rows[0].cells
        hdr[0].text = "Citimi"
        hdr[1].text = "Kodi"
        hdr[2].text = "Statusi"
        for it in items:
            row = cit_table.add_row().cells
            row[0].text = str(it.get("raw") or "")
            row[1].text = str(it.get("code_label") or "—")
            status = str(it.get("status") or "")
            row[2].text = {
                "verified": "✓ verifikuar",
                "fake": "✗ fantazmë",
                "needs_code": "? pa kod",
            }.get(status, status)

    retrieved = pack.get("retrieved_articles") or []
    if retrieved:
        doc.add_heading("Burime të marra (BM25)", level=2)
        for art in retrieved[:20]:
            line = (
                f"Neni {art.get('number') or '?'} ({art.get('code') or '—'})"
                f" — {(art.get('heading') or '')[:80]}"
            )
            score = art.get("score")
            if score is not None:
                line += f"   [score: {score}]"
            doc.add_paragraph(line, style="List Bullet")

    doc.add_paragraph()
    foot = doc.add_paragraph()
    fr = foot.add_run(
        "Ky dokument provenance gjenerohet automatikisht nga Super Avvocato "
        "për qëllime auditimi (EU AI Act art. 12-13, llogaridhënia e "
        "sistemeve me rrezik të lartë)."
    )
    fr.italic = True
    fr.font.size = Pt(8)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
